"""
File Processor Module
Handles extraction of content from various file formats
"""

import os
from pathlib import Path


class FileProcessor:
    """Processes and extracts content from various file formats"""

    def extract_content(self, filepath):
        """
        Extract content from a file based on its type
        
        Args:
            filepath: Path to the file
        
        Returns:
            Extracted content as string
        """
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"File not found: {filepath}")
        
        file_ext = Path(filepath).suffix.lower()
        
        if file_ext == '.pdf':
            return self._extract_pdf(filepath)
        elif file_ext == '.docx':
            return self._extract_docx(filepath)
        elif file_ext == '.txt':
            return self._extract_txt(filepath)
        elif file_ext in ['.xlsx', '.xls']:
            return self._extract_excel(filepath)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")

    def _extract_pdf(self, filepath):
        """Extract text from PDF file"""
        try:
            import PyPDF2
            content = []
            with open(filepath, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    content.append(page.extract_text())
            return '\n'.join(content)
        except ImportError:
            return "PyPDF2 not available. Install it to process PDF files."
        except Exception as e:
            raise Exception(f"Error extracting PDF: {str(e)}")

    def _extract_docx(self, filepath):
        """Extract text from DOCX file"""
        try:
            from docx import Document
            doc = Document(filepath)
            content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            content.append(cell.text)
            
            return '\n'.join(content)
        except ImportError:
            return "python-docx not available. Install it to process DOCX files."
        except Exception as e:
            raise Exception(f"Error extracting DOCX: {str(e)}")

    def _extract_txt(self, filepath):
        """Extract text from TXT file"""
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try different encoding
            with open(filepath, 'r', encoding='latin-1') as f:
                return f.read()
        except Exception as e:
            raise Exception(f"Error extracting TXT: {str(e)}")

    def _extract_excel(self, filepath):
        """Extract data from Excel file"""
        try:
            import openpyxl
            content = []
            workbook = openpyxl.load_workbook(filepath)
            
            for sheet in workbook.sheetnames:
                ws = workbook[sheet]
                content.append(f"\n=== Sheet: {sheet} ===\n")
                
                for row in ws.iter_rows(values_only=True):
                    row_content = [str(cell) if cell is not None else "" for cell in row]
                    content.append(' | '.join(row_content))
            
            return '\n'.join(content)
        except ImportError:
            return "openpyxl not available. Install it to process Excel files."
        except Exception as e:
            raise Exception(f"Error extracting Excel: {str(e)}")

    def validate_file(self, filepath, allowed_extensions=None):
        """
        Validate if file is acceptable
        
        Args:
            filepath: Path to the file
            allowed_extensions: List of allowed extensions (default: common business formats)
        
        Returns:
            Tuple (is_valid, error_message)
        """
        if allowed_extensions is None:
            allowed_extensions = ['pdf', 'docx', 'txt', 'xlsx', 'xls']
        
        if not os.path.exists(filepath):
            return False, "File does not exist"
        
        file_ext = Path(filepath).suffix.lower().lstrip('.')
        if file_ext not in allowed_extensions:
            return False, f"File type .{file_ext} is not supported"
        
        # Check file size (max 50MB)
        file_size = os.path.getsize(filepath)
        if file_size > 50 * 1024 * 1024:
            return False, "File size exceeds 50MB limit"
        
        return True, "Valid"
