"""
Case Study Generator Module
Generates templated case studies using OpenAI API
"""

import json
import openai
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


class CaseStudyGenerator:
    """Generates case studies from project content"""

    def __init__(self):
        self.model = "gpt-4"
        self.template = self._get_template()

    def _get_template(self):
        """Get case study template structure"""
        return {
            "problem_statement": "Concise description of the business challenge",
            "solution_approach": "Overview of the solution implemented",
            "key_metrics": ["metric1", "metric2", "metric3"],
            "impact_summary": "Quantifiable business impact",
            "implementation_details": "Key steps in the solution",
            "client_testimonial": "Optional client feedback",
            "lessons_learned": "Key insights and learnings"
        }

    def generate(self, project_name, client_name, industry, extracted_content, additional_context=""):
        """
        Generate a case study from project content
        
        Args:
            project_name: Name of the project
            client_name: Name of the client
            industry: Industry vertical
            extracted_content: Dictionary of extracted file contents
            additional_context: Additional context provided by user
        
        Returns:
            Dictionary containing the structured case study
        """
        
        # Prepare content for LLM
        content_summary = self._prepare_content_summary(extracted_content)
        
        # Create prompt for case study generation
        prompt = self._create_prompt(
            project_name=project_name,
            client_name=client_name,
            industry=industry,
            content_summary=content_summary,
            additional_context=additional_context
        )

        # Call OpenAI API
        response = openai.ChatCompletion.create(
            model=self.model,
            messages=[
                {
                    "role": "system",
                    "content": """You are an expert business consultant specializing in creating 
                    compelling case studies. Generate a structured one-page case study that clearly 
                    articulates the problem, solution, and quantifiable impact. Use professional language 
                    and focus on business value."""
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.7,
            max_tokens=2000
        )

        # Parse response
        case_study_content = response.choices[0].message.content
        case_study = self._parse_case_study(case_study_content)
        
        # Add metadata
        case_study['metadata'] = {
            'project_name': project_name,
            'client_name': client_name,
            'industry': industry,
            'model_used': self.model
        }

        return case_study

    def _prepare_content_summary(self, extracted_content):
        """Prepare a summary of extracted content"""
        summary = ""
        for filename, content in extracted_content.items():
            summary += f"\n=== {filename} ===\n"
            # Limit content to first 1000 characters to avoid token limits
            if isinstance(content, str):
                summary += content[:1000] + ("..." if len(content) > 1000 else "")
            else:
                summary += str(content)[:1000]
        return summary

    def _create_prompt(self, project_name, client_name, industry, content_summary, additional_context):
        """Create the prompt for case study generation"""
        prompt = f"""
        Generate a one-page case study for the following project:
        
        Project Name: {project_name}
        Client Name: {client_name}
        Industry: {industry}
        
        Additional Context from User: {additional_context or 'None provided'}
        
        Project Deliverables and Documentation:
        {content_summary}
        
        Please structure the case study with the following sections:
        
        1. PROBLEM STATEMENT: What was the business challenge or opportunity?
        2. SOLUTION APPROACH: How was the problem addressed?
        3. KEY METRICS: What were the key performance indicators or metrics?
        4. IMPACT SUMMARY: What was the quantifiable business impact?
        5. IMPLEMENTATION DETAILS: What were the key steps in the implementation?
        6. LESSONS LEARNED: What insights were gained?
        
        Provide the response in valid JSON format with these exact keys:
        problem_statement, solution_approach, key_metrics (array), impact_summary, 
        implementation_details, lessons_learned
        
        Ensure all information is accurate to the provided documentation and 
        maintains a professional tone suitable for executive reading.
        """
        return prompt

    def _parse_case_study(self, content):
        """Parse case study content from LLM response"""
        try:
            # Try to extract JSON from the response
            import re
            json_match = re.search(r'\{.*\}', content, re.DOTALL)
            if json_match:
                case_study = json.loads(json_match.group())
            else:
                # If no JSON found, create structured output from text
                case_study = {
                    "problem_statement": content[:500],
                    "solution_approach": content[500:1000],
                    "key_metrics": ["Metric 1", "Metric 2", "Metric 3"],
                    "impact_summary": content[1000:1500],
                    "implementation_details": content[1500:2000],
                    "lessons_learned": "Key insights from the project implementation"
                }
        except json.JSONDecodeError:
            case_study = {
                "problem_statement": content[:500],
                "solution_approach": content[500:1000],
                "key_metrics": ["Metric 1", "Metric 2", "Metric 3"],
                "impact_summary": content[1000:1500],
                "implementation_details": content[1500:2000],
                "lessons_learned": "Key insights from the project implementation"
            }
        
        return case_study

    def save_to_docx(self, case_study, filepath):
        """Save case study to a Word document"""
        doc = Document()
        
        # Add title
        title = doc.add_paragraph()
        title_run = title.add_run(f"{case_study['metadata']['project_name']} - Case Study")
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        metadata_para = doc.add_paragraph()
        metadata_run = metadata_para.add_run(
            f"Client: {case_study['metadata']['client_name']} | Industry: {case_study['metadata']['industry']}"
        )
        metadata_run.font.size = Pt(10)
        metadata_run.font.italic = True
        metadata_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing
        
        # Add sections
        sections = [
            ("Problem Statement", case_study.get('problem_statement', '')),
            ("Solution Approach", case_study.get('solution_approach', '')),
            ("Key Metrics", case_study.get('key_metrics', [])),
            ("Impact Summary", case_study.get('impact_summary', '')),
            ("Implementation Details", case_study.get('implementation_details', '')),
            ("Lessons Learned", case_study.get('lessons_learned', ''))
        ]
        
        for section_title, section_content in sections:
            # Add section heading
            heading = doc.add_heading(section_title, level=2)
            heading.style = 'Heading 2'
            
            # Add section content
            if isinstance(section_content, list):
                for item in section_content:
                    doc.add_paragraph(str(item), style='List Bullet')
            else:
                doc.add_paragraph(str(section_content))
            
            doc.add_paragraph()  # Spacing between sections
        
        # Add footer with generation info
        footer_para = doc.add_paragraph()
        footer_para.paragraph_format.left_indent = Inches(0)
        footer_run = footer_para.add_run(
            f"\nGenerated using Case Study AI Generator | {case_study['metadata']['model_used']}"
        )
        footer_run.font.size = Pt(8)
        footer_run.font.italic = True
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Save document
        doc.save(filepath)
        return filepath

    def save_to_json(self, case_study, filepath):
        """Save case study to JSON file"""
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(case_study, f, indent=2, ensure_ascii=False)
        return filepath
